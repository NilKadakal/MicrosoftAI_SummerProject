import json
import sqlite3
from pathlib import Path

from docx import Document
from foundry_local_sdk import Configuration, FoundryLocalManager
import pdfplumber


DOCS_DIR = Path("data/docs")
DB_PATH = Path("rag.db")
SUPPORTED_EXTENSIONS = {".txt", ".pdf", ".docx"}


def read_txt(path):
    return path.read_text(encoding="utf-8")


def format_table(table):
    lines = []

    for row in table:
        cleaned_cells = []

        for cell in row:
            if cell is None:
                cleaned_cells.append("")
            else:
                cleaned_cells.append(str(cell).strip())

        line = " | ".join(cleaned_cells)

        if line.strip():
            lines.append(line)

    return "\n".join(lines)


def read_pdf(path):
    parts = []

    with pdfplumber.open(path) as pdf:
        for page_number, page in enumerate(pdf.pages, start=1):
            text = page.extract_text()

            if text:
                parts.append(f"Page {page_number} text:\n{text.strip()}")

            tables = page.extract_tables()

            for table_index, table in enumerate(tables, start=1):
                table_text = format_table(table)

                if table_text:
                    parts.append(
                        f"Page {page_number} table {table_index}:\n{table_text}"
                    )

    return "\n\n".join(parts)


def read_docx(path):
    document = Document(path)
    parts = []

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            parts.append(text)

    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))

    return "\n\n".join(parts)


def read_document(path):
    extension = path.suffix.lower()

    if extension == ".txt":
        return read_txt(path)

    if extension == ".pdf":
        return read_pdf(path)

    if extension == ".docx":
        return read_docx(path)

    raise ValueError(f"Unsupported file type: {path.name}")


def chunk_text(text, paragraphs_per_chunk=2):
    paragraphs = [
        paragraph.strip()
        for paragraph in text.split("\n\n")
        if paragraph.strip()
    ]

    chunks = []

    for i in range(0, len(paragraphs), paragraphs_per_chunk):
        chunk = "\n\n".join(paragraphs[i:i + paragraphs_per_chunk])
        chunks.append(chunk)

    return chunks


def setup_database():
    connection = sqlite3.connect(DB_PATH)
    cursor = connection.cursor()

    cursor.execute("""
        CREATE TABLE IF NOT EXISTS chunks (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            source TEXT NOT NULL,
            chunk_index INTEGER NOT NULL,
            chunk_text TEXT NOT NULL,
            embedding TEXT NOT NULL
        )
    """)

    cursor.execute("DELETE FROM chunks")
    connection.commit()

    return connection


def save_chunk(connection, source, chunk_index, chunk_text_value, embedding):
    cursor = connection.cursor()

    cursor.execute(
        """
        INSERT INTO chunks (source, chunk_index, chunk_text, embedding)
        VALUES (?, ?, ?, ?)
        """,
        (source, chunk_index, chunk_text_value, json.dumps(embedding))
    )

    connection.commit()


def get_document_paths():
    return sorted(
        path
        for path in DOCS_DIR.iterdir()
        if path.is_file() and path.suffix.lower() in SUPPORTED_EXTENSIONS
    )


def main():
    if not DOCS_DIR.exists():
        print("data/docs folder was not found.")
        return

    document_paths = get_document_paths()

    if not document_paths:
        print("No supported documents found in data/docs.")
        print("Supported file types: .txt, .pdf, .docx")
        return

    config = Configuration(app_name="rag_ingestion")
    FoundryLocalManager.initialize(config)
    manager = FoundryLocalManager.instance

    print("Loading embedding model...")
    embedding_model = manager.catalog.get_model("qwen3-embedding-0.6b")
    embedding_model.download(
        lambda p: print(f"\rDownloading embedding model: {p:.1f}%", end="", flush=True)
    )
    print()
    embedding_model.load()

    embedding_client = embedding_model.get_embedding_client()

    connection = setup_database()

    total_chunks = 0

    for document_path in document_paths:
        print(f"\nReading {document_path.name}...")

        text = read_document(document_path)

        if not text.strip():
            print(f"Skipped {document_path.name}: no readable text found.")
            continue

        chunks = chunk_text(text)

        print(f"Processing {document_path.name}: {len(chunks)} chunks")

        for chunk_index, chunk in enumerate(chunks):
            response = embedding_client.generate_embedding(chunk)
            embedding = response.data[0].embedding

            save_chunk(
                connection,
                source=document_path.name,
                chunk_index=chunk_index,
                chunk_text_value=chunk,
                embedding=embedding,
            )

            total_chunks += 1

    connection.close()
    embedding_model.unload()

    print(f"\nSaved {total_chunks} chunks to {DB_PATH}.")


if __name__ == "__main__":
    main()